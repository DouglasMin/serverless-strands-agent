import base64
import io
import json
import re
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from strands import tool


def _set_cell_background(cell: Any, fill_hex: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_styled_runs_to_paragraph(p: Any, text: str, default_font_size: Pt = Pt(11), default_color: RGBColor = RGBColor(0x1E, 0x29, 0x3B)):
    """Parse inline **bold** and *italic* markdown spans and append runs to paragraph."""
    # Split by bold **...**
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            clean = part[2:-2]
            r = p.add_run(clean)
            r.font.name = "Calibri"
            r.font.size = default_font_size
            r.font.bold = True
            r.font.color.rgb = default_color
        else:
            # Check for *italic*
            subparts = re.split(r"(\*.*?\*)", part)
            for sub in subparts:
                if not sub:
                    continue
                if sub.startswith("*") and sub.endswith("*") and len(sub) >= 2:
                    r = p.add_run(sub[1:-1])
                    r.font.name = "Calibri"
                    r.font.size = default_font_size
                    r.font.italic = True
                    r.font.color.rgb = default_color
                else:
                    r = p.add_run(sub)
                    r.font.name = "Calibri"
                    r.font.size = default_font_size
                    r.font.color.rgb = default_color


def _parse_markdown_to_sections(md_text: str) -> list[dict[str, Any]]:
    """Parse a markdown research dossier into structured Word sections."""
    sections: list[dict[str, Any]] = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("# "):
            sections.append({"type": "heading_1", "text": line[2:].strip()})
            i += 1
        elif line.startswith("## "):
            sections.append({"type": "heading_1", "text": line[3:].strip()})
            i += 1
        elif line.startswith("### ") or line.startswith("#### "):
            clean = re.sub(r"^#+\s*", "", line)
            sections.append({"type": "heading_2", "text": clean.strip()})
            i += 1
        elif line.startswith("> "):
            sections.append({"type": "callout", "text": line[2:].strip()})
            i += 1
        elif line.startswith("- ") or line.startswith("* ") or re.match(r"^\d+\.\s", line):
            items = []
            while i < len(lines):
                bline = lines[i].strip()
                if bline.startswith("- ") or bline.startswith("* "):
                    items.append(bline[2:].strip())
                    i += 1
                elif re.match(r"^\d+\.\s", bline):
                    items.append(re.sub(r"^\d+\.\s*", "", bline).strip())
                    i += 1
                elif not bline:
                    i += 1
                    break
                else:
                    break
            if items:
                sections.append({"type": "bullet_list", "items": items})
        elif line.startswith("|") and "|" in line[1:]:
            # Parse markdown table
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                data_start = 1
                if "---" in table_lines[1] or ":---" in table_lines[1]:
                    data_start = 2
                rows = []
                for tline in table_lines[data_start:]:
                    cells = [c.strip() for c in tline.split("|")[1:-1]]
                    rows.append(cells)
                sections.append({"type": "table", "headers": headers, "rows": rows})
        else:
            para_lines = []
            while i < len(lines):
                pline = lines[i].strip()
                if (
                    not pline
                    or pline.startswith("#")
                    or pline.startswith("|")
                    or pline.startswith("- ")
                    or pline.startswith("* ")
                    or pline.startswith("> ")
                    or re.match(r"^\d+\.\s", pline)
                ):
                    break
                para_lines.append(pline)
                i += 1
            if para_lines:
                sections.append({"type": "paragraph", "text": " ".join(para_lines)})

    return sections


@tool
def create_word_document(
    filename: str,
    title: str,
    sections_json: str = "",
    content: str = "",
) -> str:
    """Create a formatted Microsoft Word (.docx) document with styled headings, paragraphs, bullet lists, tables, and callouts.

    Args:
        filename: Name of the file, e.g. "Humanoid_Robotics_Report.docx"
        title: Document main title
        sections_json: Optional JSON string representing structured sections.
        content: Optional raw Markdown text or report body (automatically parsed into headings, tables, bullets, and paragraphs).
    """
    if not filename.endswith(".docx"):
        filename += ".docx"

    sections: list[dict[str, Any]] = []

    # 1. Try parsing sections_json
    if sections_json and sections_json.strip():
        try:
            parsed = json.loads(sections_json)
            if isinstance(parsed, list):
                sections = parsed
            elif isinstance(parsed, dict):
                sections = [parsed]
        except Exception:
            # If JSON parsing failed, treat sections_json as markdown
            sections = _parse_markdown_to_sections(sections_json)

    # 2. If no sections from JSON, try markdown content
    if not sections and content and content.strip():
        sections = _parse_markdown_to_sections(content)

    if not sections:
        sections = [
            {"type": "heading_1", "text": title},
            {"type": "paragraph", "text": "Document generated successfully."},
        ]

    doc = Document()

    # Set page margins to 1 inch
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # Document Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    title_p.paragraph_format.space_after = Pt(8)

    for s in sections:
        stype = s.get("type", "paragraph")
        text = str(s.get("text", ""))

        if stype == "subtitle":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(13)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            p.paragraph_format.space_after = Pt(16)

        elif stype == "heading_1":
            h = doc.add_paragraph()
            r = h.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(16)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(6)

        elif stype == "heading_2":
            h = doc.add_paragraph()
            r = h.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)

        elif stype == "paragraph":
            p = doc.add_paragraph()
            _add_styled_runs_to_paragraph(p, text, default_font_size=Pt(11))
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)

        elif stype == "bullet_list":
            items = s.get("items", [])
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                _add_styled_runs_to_paragraph(p, str(item), default_font_size=Pt(11))
                p.paragraph_format.space_after = Pt(3)

        elif stype == "callout":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            r_icon = p.add_run("💡 ")
            r_icon.font.size = Pt(11)
            r = p.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(10.5)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

        elif stype == "table":
            headers = s.get("headers", [])
            rows = s.get("rows", [])
            if headers:
                table = doc.add_table(
                    rows=len(rows) + 1, cols=len(headers)
                )
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True

                # Style Header Row
                for col_idx, h_text in enumerate(headers):
                    cell = table.cell(0, col_idx)
                    cell.text = str(h_text)
                    _set_cell_background(cell, "1E3A8A")
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(10.5)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                # Style Data Rows
                for row_idx, row_data in enumerate(rows, 1):
                    is_alt = (row_idx % 2) == 0
                    for col_idx, val in enumerate(row_data):
                        if col_idx < len(headers):
                            cell = table.cell(row_idx, col_idx)
                            cell.text = str(val)
                            if is_alt:
                                _set_cell_background(cell, "F8FAFC")
                            p = cell.paragraphs[0]
                            for run in p.runs:
                                run.font.name = "Calibri"
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

                doc.add_paragraph().paragraph_format.space_after = Pt(6)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    size_bytes = len(docx_bytes)
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    from .s3_storage import upload_deliverable_to_s3

    s3_uri, download_url, fallback_data_uri = upload_deliverable_to_s3(
        file_bytes=docx_bytes,
        filename=filename,
        content_type=content_type,
    )

    doc_event = {
        "filename": filename,
        "file_type": "word",
        "size_bytes": size_bytes,
        "s3_uri": s3_uri,
        "url": download_url,
        **({"data_uri": fallback_data_uri} if fallback_data_uri else {}),
        "summary": f"{title} ({len(sections)} sections)",
        "title": title,
        "sections": sections,
    }

    from office_tools import document_queue

    document_queue.put_nowait(doc_event)

    return json.dumps(
        {
            "status": "success",
            "filename": filename,
            "file_type": "word",
            "size_bytes": size_bytes,
            "title": title,
            "s3_uri": s3_uri,
            "url": download_url,
            "download_ready": True,
            "summary": f"Successfully generated Word document '{filename}' with {len(sections)} sections.",
        },
        indent=2,
    )
